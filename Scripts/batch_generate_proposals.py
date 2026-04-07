#!/usr/bin/env python3
"""
Batch Proposal Generator for AI Local Website Agency

Reads lead data from Lead-Tracker_v1.0.xlsx and generates personalized
DOCX proposal documents for each warm lead.

Usage:
    python3 batch_generate_proposals.py [--output-dir path/to/output]
"""

import sys
import os
from datetime import datetime
from pathlib import Path
import re

try:
    import openpyxl
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("✗ Missing dependencies. Install with:")
    print("  pip install openpyxl python-docx")
    sys.exit(1)


class ProposalGenerator:
    def __init__(self, template_path, output_dir=None):
        self.template_path = template_path
        self.output_dir = output_dir or Path(template_path).parent / "Generated"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
    
    def load_leads(self, excel_path):
        """Load lead data from Lead-Tracker spreadsheet."""
        try:
            workbook = openpyxl.load_workbook(excel_path)
            # Assume first sheet contains lead data
            sheet = workbook.active
            
            leads = []
            headers = None
            
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                if i == 0:
                    headers = row
                    continue
                
                if not any(row):  # Skip empty rows
                    continue
                
                lead_data = dict(zip(headers, row))
                if lead_data.get('Status') == 'Warm':  # Filter for warm leads
                    leads.append(lead_data)
            
            return leads
        except Exception as e:
            print(f"✗ Error reading Excel: {e}")
            return []
    
    def sanitize_filename(self, name):
        """Convert business name to safe filename."""
        return re.sub(r'[^\w\s-]', '', name).replace(' ', '_')[:50]
    
    def personalize_document(self, doc, lead_data):
        """Replace placeholders in document with lead-specific data."""
        business_name = lead_data.get('Business Name', 'Your Business')
        demo_url = lead_data.get('Demo URL', 'https://preview.ailocal.com/[business-name]')
        contact_email = lead_data.get('Contact Email', 'hello@ailocal.com')
        contact_phone = lead_data.get('Phone', '(888) 555-0123')
        
        # Replace placeholders in all paragraphs
        for paragraph in doc.paragraphs:
            if '[BUSINESS_NAME]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[BUSINESS_NAME]', business_name)
            if '[DEMO_SITE_URL]' in paragraph.text:
                paragraph.text = paragraph.text.replace('[DEMO_SITE_URL]', demo_url)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '[BUSINESS_NAME]' in paragraph.text:
                            paragraph.text = paragraph.text.replace('[BUSINESS_NAME]', business_name)
        
        # Update footer with contact info
        footer = doc.sections[0].footer
        for paragraph in footer.paragraphs:
            text = paragraph.text
            text = text.replace('(888) 555-0123', contact_phone)
            text = text.replace('hello@ailocal.com', contact_email)
            paragraph.text = text
        
        return doc
    
    def generate_proposal(self, lead_data):
        """Generate personalized proposal for a single lead."""
        try:
            doc = Document(self.template_path)
            doc = self.personalize_document(doc, lead_data)
            
            business_name = lead_data.get('Business Name', 'Unknown')
            safe_name = self.sanitize_filename(business_name)
            date_str = datetime.now().strftime('%Y%m%d')
            filename = f"{safe_name}_Proposal_{date_str}.docx"
            
            output_path = self.output_dir / filename
            doc.save(output_path)
            
            return {
                'success': True,
                'lead': business_name,
                'path': str(output_path),
                'size': os.path.getsize(output_path)
            }
        except Exception as e:
            return {
                'success': False,
                'lead': lead_data.get('Business Name', 'Unknown'),
                'error': str(e)
            }
    
    def generate_batch(self, lead_data_list):
        """Generate proposals for multiple leads."""
        results = []
        for lead_data in lead_data_list:
            result = self.generate_proposal(lead_data)
            results.append(result)
        
        return results


def main():
    # Paths
    template_path = Path('/sessions/peaceful-jolly-ritchie/mnt/Website Agency/Templates/Proposals/Client_Proposal_v1.0.docx')
    tracker_path = Path('/sessions/peaceful-jolly-ritchie/mnt/Website Agency/Templates/Lead-Tracker_v1.0.xlsx')
    output_dir = Path('/sessions/peaceful-jolly-ritchie/mnt/Website Agency/Outputs/Generated-Proposals')
    
    print("=" * 70)
    print("BATCH PROPOSAL GENERATOR")
    print("=" * 70)
    
    # Check files exist
    if not template_path.exists():
        print(f"✗ Template not found: {template_path}")
        return
    
    if not tracker_path.exists():
        print(f"✗ Lead tracker not found: {tracker_path}")
        print("  Skipping batch generation (no leads data available)")
        return
    
    # Initialize generator
    try:
        generator = ProposalGenerator(str(template_path), output_dir)
        print(f"\n✓ Generator initialized")
        print(f"  Template: {template_path.name}")
        print(f"  Output directory: {output_dir}")
    except Exception as e:
        print(f"✗ Failed to initialize generator: {e}")
        return
    
    # Load leads
    print(f"\n✓ Loading leads from: {tracker_path.name}")
    leads = generator.load_leads(str(tracker_path))
    
    if not leads:
        print("  ℹ No warm leads found in tracker (or Excel file is empty)")
        return
    
    print(f"  Found {len(leads)} warm lead(s)")
    
    # Generate proposals
    print(f"\n✓ Generating proposals...")
    results = generator.generate_batch(leads)
    
    # Report results
    successful = [r for r in results if r['success']]
    failed = [r for r in results if not r['success']]
    
    print(f"\n✓ Results:")
    print(f"  ✓ Successful: {len(successful)}")
    for r in successful:
        size_kb = r['size'] / 1024
        print(f"    - {r['lead']}: {r['path'].split('/')[-1]} ({size_kb:.1f} KB)")
    
    if failed:
        print(f"  ✗ Failed: {len(failed)}")
        for r in failed:
            print(f"    - {r['lead']}: {r['error']}")
    
    print("\n" + "=" * 70)
    print(f"✓ BATCH GENERATION COMPLETE")
    print("=" * 70)


if __name__ == '__main__':
    main()
