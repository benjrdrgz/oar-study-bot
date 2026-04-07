#!/usr/bin/env python3
"""
Fill in Service Agreement contract placeholders with client-specific information.
Usage: python3 fill_contract.py --agency "Your Agency" --client "John Smith" --business "Smith Plumbing" --date "2026-03-25" --state "Wisconsin" --billing-day "15"
"""

import argparse
import sys
from docx import Document
from pathlib import Path

def fill_contract(agency, client, business, date, state, billing_day, output_name=None):
    """Fill in contract placeholders."""

    # Load template
    template_path = Path(__file__).parent.parent / 'Templates' / 'Contracts' / 'Service_Agreement_v1.0.docx'

    if not template_path.exists():
        print(f'✗ Template not found: {template_path}')
        return False

    try:
        doc = Document(str(template_path))
    except Exception as e:
        print(f'✗ Failed to load template: {e}')
        return False

    # Define replacements
    replacements = {
        '[AGENCY_NAME]': agency,
        '[CLIENT_NAME]': client,
        '[CLIENT_BUSINESS_NAME]': business,
        '[DATE]': date,
        '[STATE]': state,
        '[BILLING_DAY]': billing_day,
    }

    # Replace in all paragraphs
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)
                for run in para.runs:
                    for old_key, new_val in replacements.items():
                        if old_key in run.text:
                            run.text = run.text.replace(old_key, new_val)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            para.text = para.text.replace(old, new)
                            for run in para.runs:
                                if old in run.text:
                                    run.text = run.text.replace(old, new)

    # Determine output path
    if output_name is None:
        output_name = f'Service_Agreement_{business.replace(" ", "_")}_v1.0.docx'

    output_path = template_path.parent / output_name

    # Save
    try:
        doc.save(str(output_path))
        print(f'✓ Contract filled and saved: {output_path}')
        print(f'✓ Agency: {agency}')
        print(f'✓ Client: {client} ({business})')
        print(f'✓ Effective Date: {date}')
        print(f'✓ Governing Law: {state}')
        print(f'✓ Billing Day: {billing_day}th of month')
        return True
    except Exception as e:
        print(f'✗ Failed to save contract: {e}')
        return False

def main():
    parser = argparse.ArgumentParser(
        description='Fill in Service Agreement contract placeholders'
    )
    parser.add_argument('--agency', required=True, help='Agency name')
    parser.add_argument('--client', required=True, help='Client name')
    parser.add_argument('--business', required=True, help='Client business name')
    parser.add_argument('--date', required=True, help='Effective date (e.g., 2026-03-25)')
    parser.add_argument('--state', required=True, help='State for governing law')
    parser.add_argument('--billing-day', required=True, help='Billing day of month (e.g., 15)')
    parser.add_argument('--output', help='Output filename (optional)')

    args = parser.parse_args()

    success = fill_contract(
        agency=args.agency,
        client=args.client,
        business=args.business,
        date=args.date,
        state=args.state,
        billing_day=args.billing_day,
        output_name=args.output
    )

    return 0 if success else 1

if __name__ == '__main__':
    sys.exit(main())
