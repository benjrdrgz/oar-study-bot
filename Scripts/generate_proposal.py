from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, RGBColor
import os

# Create document
doc = Document()

# Set up margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Colors
DARK_BLUE = RGBColor(31, 71, 136)
ACCENT_BLUE = RGBColor(74, 144, 226)
LIGHT_GRAY = RGBColor(245, 245, 245)

# Helper to add footer
def add_footer(doc, page_num):
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Your AI Local Website Agency | (888) 555-0123 | hello@ailocal.com | Page {page_num}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(102, 102, 102)
        run.font.italic = True
    
    # Add border to footer
    pPr = footer_para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'CCCCCC')
    pBdr.append(top)
    pPr.append(pBdr)

def shade_cell(cell, color):
    """Add shading to a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_vertical_align(cell, align):
    """Set vertical alignment in a cell"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:vAlign')
    tcVAlign.set(qn('w:val'), align)
    tcPr.append(tcVAlign)

# PAGE 1: COVER
doc.add_paragraph()  # Spacing
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Your New Website is Ready\n— Here's What's Next")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = DARK_BLUE

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("A Custom Website Proposal for [BUSINESS_NAME]")
run.font.size = Pt(16)
run.font.italic = True
run.font.color.rgb = ACCENT_BLUE

doc.add_paragraph()  # Spacing

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Prepared by: Your AI Local Website Agency\nDate: March 25, 2026")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_paragraph()  # Spacing
preview_title = doc.add_paragraph()
preview_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = preview_title.add_run("View Your Preview Site\n")
run.font.size = Pt(12)
run.font.bold = True
run.font.color.rgb = ACCENT_BLUE

preview_url = doc.add_paragraph()
preview_url.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = preview_url.add_run("[DEMO_SITE_URL]")
run.font.size = Pt(11)
run.font.underline = True
run.font.color.rgb = ACCENT_BLUE

doc.add_paragraph()  # Spacing

intro1 = doc.add_paragraph("We noticed [BUSINESS_NAME] doesn't have a professional online presence. In today's digital-first world, that's leaving money on the table.")
intro1.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in intro1.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(51, 51, 51)

intro2 = doc.add_paragraph("We built a professional, custom website tailored to your business and uploaded it to the web. You can see it by clicking the preview link above. It's fully functional, mobile-friendly, and optimized for Google.")
intro2.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in intro2.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(51, 51, 51)

intro3 = doc.add_paragraph("Here's what it takes to go live:")
intro3.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in intro3.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 51, 51)

# Bullet list
bullets = [
    "Choose your service plan (see inside for options)",
    "Sign the simple one-page agreement",
    "We activate your website within 24 hours",
    "Customers start finding you on Google"
]
for bullet in bullets:
    p = doc.add_paragraph(bullet, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(51, 51, 51)

closing = doc.add_paragraph("No long-term contracts. No setup hassles. No surprises.")
closing.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in closing.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = ACCENT_BLUE

# Page break
doc.add_page_break()

# PAGE 2: WHAT YOU GET
title2 = doc.add_paragraph("What You Get")
title2.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in title2.runs:
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

features = [
    "Professional, custom website — built for your industry",
    "Mobile-responsive design — works perfect on phones & tablets",
    "Fast loading speed — optimized for Google ranking",
    "SSL security certificate — the lock icon customers trust",
    "Professional email setup — [yourname@yourbusiness.com]",
    "Hosting and maintenance — we handle all the technical stuff",
    "Monthly updates as needed — keep your site fresh",
    "Google optimization — customers can find you when they search"
]

for feature in features:
    p = doc.add_paragraph(feature, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(51, 51, 51)

# Why This Matters
title3 = doc.add_paragraph("\nWhy This Matters")
title3.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in title3.runs:
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

stats_intro = doc.add_paragraph("The numbers tell the story:")
for run in stats_intro.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(51, 51, 51)

# Stats table
stats_table = doc.add_table(rows=3, cols=1)
stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER

stats_data = [
    ("97%", "of consumers search online for local businesses before visiting"),
    ("75%", "judge a business's credibility by its website design & professionalism"),
    ("2-3x", "more customer inquiries from businesses with professional websites")
]

for i, (stat, text) in enumerate(stats_data):
    row = stats_table.rows[i]
    cell = row.cells[0]
    set_cell_vertical_align(cell, 'center')
    
    if i % 2 == 0:
        shade_cell(cell, "F5F5F5")
    
    # Add stat number
    p_stat = cell.paragraphs[0]
    run = p_stat.add_run(stat)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = ACCENT_BLUE
    
    # Add stat text
    p_text = cell.add_paragraph(text)
    for run in p_text.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(51, 51, 51)

closing2 = doc.add_paragraph("\nYour website is not a luxury. It's table stakes.")
closing2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in closing2.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = ACCENT_BLUE

# Page break
doc.add_page_break()

# PAGE 3: PRICING
title4 = doc.add_paragraph("Simple Pricing. No Surprises.")
title4.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in title4.runs:
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

subtitle4 = doc.add_paragraph("Choose the plan that fits your business. All include hosting, maintenance, and 24-hour activation.")
for run in subtitle4.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)

# Pricing table
pricing_table = doc.add_table(rows=13, cols=4)
pricing_table.style = 'Light Grid Accent 1'

# Header row
header_cells = pricing_table.rows[0].cells
headers = ['Feature', 'Quick Start', 'Growth\n★ Most Popular ★', 'Dominate']
for i, header_text in enumerate(headers):
    cell = header_cells[i]
    shade_cell(cell, '1F4788')
    if i > 1:
        shade_cell(cell, '4A90E2' if i == 2 else '1F4788')
    
    cell.text = header_text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            if i == 2:
                run.font.color.rgb = RGBColor(255, 215, 0)

# Data rows
pricing_data = [
    ('Setup Fee', '$997', '$1,997', '$3,497'),
    ('Monthly Fee', '$149/mo', '$249/mo', '$397/mo'),
    ('Professional Website', '✓', '✓', '✓'),
    ('Mobile Optimized', '✓', '✓', '✓'),
    ('SSL Certificate (Secure)', '✓', '✓', '✓'),
    ('Hosting & Maintenance', '✓', '✓', '✓'),
    ('Monthly Updates', '1', '3', 'Unlimited'),
    ('Google Business Setup', '—', '✓', '✓'),
    ('Directory Listings', '—', '3', '10+'),
    ('SEO Package', 'Basic', 'Standard', 'Premium'),
    ('Review Management', '—', '—', '✓'),
    ('Priority Support', '—', '—', '✓'),
]

for row_idx, row_data in enumerate(pricing_data, 1):
    row = pricing_table.rows[row_idx]
    for col_idx, cell_text in enumerate(row_data):
        cell = row.cells[col_idx]
        
        # Alternating row colors
        if row_idx % 2 == 0:
            shade_cell(cell, 'F5F5F5' if col_idx != 2 else '4A90E2')
        else:
            shade_cell(cell, 'FFFFFF' if col_idx != 2 else '4A90E2')
        
        # Special coloring for header fees
        if row_idx <= 2 and col_idx > 0:
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.bold = True
                    if col_idx == 2:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        run.font.color.rgb = RGBColor(51, 51, 51)
        else:
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                if col_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    if col_idx == 2 and cell_text == '✓':
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    elif cell_text == '✓':
                        run.font.color.rgb = ACCENT_BLUE
                    elif col_idx == 2:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        run.font.color.rgb = RGBColor(51, 51, 51)

footnote = doc.add_paragraph("All plans include your own custom domain name and professional email.")
footnote.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footnote.runs:
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(102, 102, 102)

# Page break
doc.add_page_break()

# PAGE 4: NEXT STEPS
title5 = doc.add_paragraph("Getting Started is Simple")
title5.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in title5.runs:
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

# Steps table
steps_table = doc.add_table(rows=3, cols=2)
steps_table.alignment = WD_TABLE_ALIGNMENT.LEFT

steps_data = [
    ('1', 'Choose Your Plan', 'Pick the tier that fits your business and budget.'),
    ('2', 'We Go Live in 24 Hours', 'Sign the agreement and we activate your site immediately.'),
    ('3', 'Start Getting Found on Google', 'Customers find you when they search. Inquiries start coming in.')
]

for row_idx, (num, title, desc) in enumerate(steps_data):
    row = steps_table.rows[row_idx]
    
    # Number cell
    num_cell = row.cells[0]
    shade_cell(num_cell, '4A90E2' if row_idx % 2 == 0 else 'F5F5F5')
    num_cell.text = num
    for paragraph in num_cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255) if row_idx % 2 == 0 else ACCENT_BLUE
    
    # Content cell
    content_cell = row.cells[1]
    shade_cell(content_cell, 'F5F5F5' if row_idx % 2 == 0 else 'FFFFFF')
    content_cell.text = ''
    
    p_title = content_cell.paragraphs[0]
    run = p_title.add_run(title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    
    p_desc = content_cell.add_paragraph(desc)
    for run in p_desc.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(102, 102, 102)

# Guarantee section
doc.add_paragraph()
guarantee_title = doc.add_paragraph("Our Guarantee")
for run in guarantee_title.runs:
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

guarantee_text = doc.add_paragraph("If you're not thrilled with your website within 30 days, we'll rebuild it from scratch at no additional cost. No questions asked.")
for run in guarantee_text.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

# Add left border to guarantee
p = guarantee_text._element
pPr = p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
left = OxmlElement('w:left')
left.set(qn('w:val'), 'single')
left.set(qn('w:sz'), '24')
left.set(qn('w:space'), '24')
left.set(qn('w:color'), '4A90E2')
pBdr.append(left)
pPr.append(pBdr)

# Indent the guarantee text
p_format = guarantee_text.paragraph_format
p_format.left_indent = Inches(0.3)

# Call to action
doc.add_paragraph()
cta = doc.add_paragraph("Ready to go live?")
for run in cta.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 51, 51)

cta2 = doc.add_paragraph("Reply to this email or call us at (888) 555-0123 to get started.")
for run in cta2.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(51, 51, 51)

# Closing
doc.add_paragraph()
closing_msg = doc.add_paragraph("Looking forward to building your web presence.")
for run in closing_msg.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)

signature = doc.add_paragraph("— Benjamin Rodriguez")
for run in signature.runs:
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

title_sig = doc.add_paragraph("Chief, Your AI Local Website Agency")
for run in title_sig.runs:
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(102, 102, 102)

# Save document
output_path = '/sessions/peaceful-jolly-ritchie/mnt/Website Agency/Templates/Proposals/Client_Proposal_v1.0.docx'

# Ensure directory exists
os.makedirs(os.path.dirname(output_path), exist_ok=True)

doc.save(output_path)
print(f"✓ Document created successfully: {output_path}")
print(f"✓ File size: {os.path.getsize(output_path)} bytes")

