#!/usr/bin/env python3
"""
Generate a professional client onboarding welcome packet for AI Local Website Agency
Uses python-docx library
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# Brand colors (RGB)
DARK_BLUE = RGBColor(0, 51, 102)
BRIGHT_TEAL = RGBColor(0, 168, 204)
DARK_CHARCOAL = RGBColor(51, 51, 51)
SUCCESS_GREEN = RGBColor(39, 174, 96)
LIGHT_GRAY = RGBColor(236, 239, 241)

# Placeholder variables
AGENCY_NAME = "AI Local Website Agency"
CLIENT_NAME = "[CLIENT_NAME]"
CLIENT_BUSINESS_NAME = "[YOUR_BUSINESS_NAME]"
CONTACT_NAME = "Benjamin Rodriguez"
CONTACT_EMAIL = "benjamin@ailocalsites.com"
CONTACT_PHONE = "(555) 123-4567"

def set_cell_background(cell, fill_color):
    """Set the background color of a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading_custom(doc, text, level=1, color=DARK_BLUE):
    """Add a custom heading with specific formatting"""
    sizes = {1: 32, 2: 28, 3: 24}
    space_before = {1: 12, 2: 9, 3: 6}
    space_after = {1: 6, 2: 4, 3: 3}
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(space_before[level])
    para.paragraph_format.space_after = Pt(space_after[level])
    
    run = para.add_run(text)
    run.font.size = Pt(sizes[level])
    run.font.bold = True
    run.font.color.rgb = color
    run.font.name = 'Arial'
    
    return para

def add_body_text(doc, text, bold=False, color=DARK_CHARCOAL, size=11, space_after=6):
    """Add body paragraph with formatting"""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(space_after)
    
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color
        run.font.name = 'Arial'
    
    return para

def add_bullet_point(doc, text, size=11):
    """Add a bullet point"""
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Inches(0.25)
    
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = DARK_CHARCOAL
        run.font.name = 'Arial'
    
    return para

def add_checkbox_item(doc, text, size=11):
    """Add a checkbox list item"""
    para = doc.add_paragraph('☐  ' + text)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Inches(0.25)
    
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = DARK_CHARCOAL
        run.font.name = 'Arial'
    
    return para

def create_welcome_page(doc):
    """Create page 1: Welcome"""
    # Spacer
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Main heading
    main_heading = doc.add_paragraph()
    main_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_heading.paragraph_format.space_after = Pt(8)
    run = main_heading.add_run("Welcome to " + AGENCY_NAME + "!")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.name = 'Arial'
    
    # Congratulations
    congrats = doc.add_paragraph()
    congrats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    congrats.paragraph_format.space_after = Pt(12)
    run = congrats.add_run("Congratulations, " + CLIENT_NAME + "! Your website journey starts now.")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = BRIGHT_TEAL
    run.font.name = 'Arial'
    
    # Thank you
    add_body_text(doc, 
        "Thank you for choosing us to build your online presence. We're excited to partner with you and get your website live. You've made a smart investment in your business—one that will start bringing you customers right away.",
        space_after=8)
    
    # Next steps heading
    add_heading_custom(doc, "Here's what happens next — we'll have you live in 24-48 hours.", 2)
    
    add_body_text(doc,
        "Below, you'll find your onboarding checklist, timeline, and everything you need to know. We've made this as simple as possible because we know you're busy running your business.",
        space_after=12)
    
    # Contact section
    add_heading_custom(doc, "Your dedicated contact:", 3)
    add_body_text(doc, CONTACT_NAME, bold=True, color=DARK_BLUE)
    add_body_text(doc, "Email: " + CONTACT_EMAIL, space_after=4)
    add_body_text(doc, "Phone: " + CONTACT_PHONE, space_after=8)
    
    para = doc.add_paragraph("Reply to this email or call/text anytime. We're here to help.")
    para.paragraph_format.space_after = Pt(12)
    for run in para.runs:
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = BRIGHT_TEAL
        run.font.italic = True
        run.font.name = 'Arial'

def create_checklist_page(doc):
    """Create page 2: Onboarding Checklist"""
    add_heading_custom(doc, "Your Onboarding Checklist", 1)
    
    add_body_text(doc, "What we need from you:", bold=True, space_after=8)
    
    checklist_items = [
        "Business name (exactly as you want it displayed)",
        "Phone number for the website",
        "Business address",
        "Business hours (Mon–Sun)",
        "Email address for contact form",
        "Logo file (PNG or SVG preferred, optional)",
        "3–5 photos of your business/work (optional — we can use professional stock)",
        "List of services you offer (with brief descriptions)",
        "Any specific colors or style preferences",
        "Preferred domain name (e.g., www.yourbusiness.com)",
        "Social media links (if any)",
    ]
    
    for item in checklist_items:
        add_checkbox_item(doc, item)
    
    doc.add_paragraph()
    
    add_body_text(doc, "Don't have all of this? No problem!", bold=True, color=DARK_BLUE, space_after=4)
    add_body_text(doc,
        "Send us what you have and we'll work with it. We can always add more later. The goal is to get your site live fast—perfection can wait until next week.",
        space_after=8)
    
    add_body_text(doc, "How to submit:", bold=True, color=DARK_BLUE, space_after=4)
    add_body_text(doc,
        "Reply to this email with your business information. Attach files directly or paste text. We'll take it from there.",
        space_after=6)

def create_timeline_page(doc):
    """Create page 3: Timeline"""
    add_heading_custom(doc, "Your Timeline", 1)
    
    add_body_text(doc, "Here's exactly what happens and when:", space_after=12)
    
    # Create timeline table
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    headers = ['Day', 'What Happens', 'Your Action']
    
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        set_cell_background(cell, '003366')
        cell.text = header_text
        
        # Format header text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(11)
                run.font.name = 'Arial'
    
    # Data rows
    timeline_data = [
        ('Day 1', 'We begin customizing your website', 'Send us your business info'),
        ('Day 2-3', 'Your site is built and tested', 'Nothing — we are working!'),
        ('Day 4', 'Preview link sent for your review', 'Review and send feedback'),
        ('Day 5', 'We make any revisions', 'Approve final version'),
        ('Day 6', 'YOUR SITE IS LIVE!', 'Celebrate!'),
    ]
    
    for i, (day, what, action) in enumerate(timeline_data, 1):
        row_cells = table.rows[i].cells
        
        # Day column (centered)
        row_cells[0].text = day
        for para in row_cells[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Arial'
        
        # What Happens
        row_cells[1].text = what
        for para in row_cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Arial'
        
        # Your Action (centered)
        row_cells[2].text = action
        for para in row_cells[2].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Arial'
    
    doc.add_paragraph()
    
    para = doc.add_paragraph("That's it. Six days from submission to go-live. Most of the time it's even faster.")
    para.paragraph_format.space_after = Pt(8)
    for run in para.runs:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = BRIGHT_TEAL
        run.font.name = 'Arial'

def create_included_page(doc):
    """Create page 4: What's Included"""
    add_heading_custom(doc, "What's Included in Your Plan", 1)
    
    add_body_text(doc, "Your website comes with everything you need to succeed:", space_after=12)
    
    included_items = [
        "Website hosting and maintenance (always online)",
        "SSL security certificate (always protected)",
        "Content updates per month (included in your plan)",
        "Performance monitoring",
        "Priority email support",
        "Mobile-friendly design (works on all devices)",
        "Google-optimized setup",
        "You own the website forever",
    ]
    
    for item in included_items:
        add_bullet_point(doc, item)
    
    doc.add_paragraph()
    
    add_body_text(doc, "What you won't get:", bold=True, space_after=6)
    
    excluded_items = [
        "Monthly fees or surprise charges",
        "Long-term contracts or lock-in",
        "Hidden costs",
        "Pressure to 'upgrade' or add unnecessary features",
    ]
    
    for item in excluded_items:
        add_bullet_point(doc, item)
    
    doc.add_paragraph()
    
    add_body_text(doc, "What happens after launch?", bold=True, color=DARK_BLUE, space_after=4)
    add_body_text(doc,
        "Your website will be yours to keep. You can make updates yourself, ask us to update it for you, or even transfer it to another provider. No contracts. No lock-in. Just your website, forever.",
        space_after=8)

def create_faq_page(doc):
    """Create page 5: FAQ"""
    add_heading_custom(doc, "Frequently Asked Questions", 1)
    
    faq_items = [
        ("How do I request changes to my site?",
         "Simply email us or call. Send screenshots, descriptions, or notes about what you'd like changed. We'll update it and send you a preview. Most updates take 1–2 business days."),
        
        ("What happens if my site goes down?",
         "We monitor your site 24/7. If there's ever an issue, we'll fix it immediately. You'll never have to worry about it—that's our job."),
        
        ("Can I add more pages later?",
         "Absolutely. You can expand your site whenever you want. Just let us know what you need, and we'll add it."),
        
        ("How do I cancel?",
         "You can cancel anytime. No questions asked. The website is yours—you keep it forever. If you want us to stop providing updates, just let us know."),
        
        ("What if I want to move my site?",
         "You own the website. You can move it, redesign it, or hand it off to another provider anytime. We'll provide all files and access you need."),
        
        ("Do you help with Google Business Profile?",
         "We'll set up your profile and optimize it to work with your website. We'll also help you manage it going forward so customers can find you on Google Maps and Google Search."),
    ]
    
    for question, answer in faq_items:
        add_heading_custom(doc, question, 3)
        add_body_text(doc, answer, space_after=10)

def create_get_started_page(doc):
    """Create page 6: Get Started"""
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_heading_custom(doc, "Let's Get Started", 1)
    
    add_body_text(doc, "Here's all you need to do right now:", space_after=8)
    
    add_body_text(doc, "1. Look at the Onboarding Checklist on page 2", bold=True, space_after=4)
    add_body_text(doc, "2. Reply to this email with your business information", bold=True, space_after=4)
    add_body_text(doc, "3. We'll confirm receipt and start building immediately", bold=True, space_after=12)
    
    add_body_text(doc, "Questions right now?", bold=True, color=DARK_BLUE, space_after=4)
    add_body_text(doc, "Call or text Benjamin at " + CONTACT_PHONE + ". We're here to help.", space_after=16)
    
    para = doc.add_paragraph("We can't wait to help " + CLIENT_BUSINESS_NAME + " grow online!")
    para.paragraph_format.space_after = Pt(12)
    for run in para.runs:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = BRIGHT_TEAL
        run.font.name = 'Arial'
    
    # Signature
    add_body_text(doc, "— " + CONTACT_NAME, space_after=4)
    add_body_text(doc, AGENCY_NAME, bold=True, color=DARK_BLUE, space_after=0)

def main():
    """Generate the welcome packet document"""
    try:
        # Create document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Create all pages
        create_welcome_page(doc)
        doc.add_page_break()
        
        create_checklist_page(doc)
        doc.add_page_break()
        
        create_timeline_page(doc)
        doc.add_page_break()
        
        create_included_page(doc)
        doc.add_page_break()
        
        create_faq_page(doc)
        doc.add_page_break()
        
        create_get_started_page(doc)
        
        # Save document
        output_dir = '/sessions/peaceful-jolly-ritchie/mnt/Website Agency/Client-Onboarding'
        os.makedirs(output_dir, exist_ok=True)
        
        output_path = os.path.join(output_dir, 'Welcome_Packet_v1.0.docx')
        doc.save(output_path)
        
        # Get file size
        file_size = os.path.getsize(output_path) / 1024
        
        print("✓ Welcome packet generated successfully!")
        print("✓ Saved to: " + output_path)
        print("✓ File size: " + str(round(file_size, 2)) + " KB")
        print("✓ Pages: 6")
        print("✓ Ready for customization with client details")
        print("")
        print("Document includes:")
        print("  • Page 1: Welcome & Contact Info")
        print("  • Page 2: Onboarding Checklist")
        print("  • Page 3: Timeline Table")
        print("  • Page 4: What's Included")
        print("  • Page 5: FAQ")
        print("  • Page 6: Get Started")
        
    except Exception as e:
        print("Error generating document: " + str(e))
        import traceback
        traceback.print_exc()
        exit(1)

if __name__ == '__main__':
    main()
