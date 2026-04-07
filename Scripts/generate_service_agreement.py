#!/usr/bin/env python3
"""
Generate a professional Service Agreement contract for the AI Local Website Agency.
Uses python-docx to create a properly formatted DOCX file.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys

def set_cell_background(cell, fill):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def shade_paragraph(paragraph, fill):
    """Set paragraph background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    paragraph._element.get_or_add_pPr().append(shading_elm)

def create_heading(doc, text, level=1):
    """Add a heading to the document."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_body_paragraph(doc, text, indent=0, bold=False, italic=False):
    """Add a body paragraph with optional formatting."""
    p = doc.add_paragraph(text)
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent * 0.5)

    # Apply formatting to all runs
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic

    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)

    return p

def add_numbered_item(doc, number, text):
    """Add a numbered list item."""
    p = doc.add_paragraph(f'{number}.  {text}', style='List Number')
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet_point(doc, text):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p

def create_service_tiers_table(doc):
    """Create the service tiers comparison table."""
    table = doc.add_table(rows=15, cols=4)
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Feature', 'Quick Start\n($997 + $149/mo)', 'Growth\n($1,997 + $249/mo)', 'Dominate\n($3,497 + $397/mo)']

    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        set_cell_background(cell, 'D3D3D3')
        cell.text = header_text
        # Format header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)

    # Data rows
    data = [
        ['Professional Website', '5-page\nsingle-page', '5-10 page\nmulti-page', 'Unlimited pages'],
        ['Mobile Optimization', '✓ Included', '✓ Included', '✓ Included'],
        ['Hosting & Maintenance', '✓ Included', '✓ Included', '✓ Included'],
        ['SSL Certificate', '✓ Included', '✓ Included', '✓ Included'],
        ['Basic SEO Setup', '✓ Included', '✓ Included', '✓ Included'],
        ['Revision Rounds', '1 round', '3 rounds', 'Unlimited'],
        ['Google Business Profile Opt.', '—', '✓ Included', '✓ Included'],
        ['Directory Listings', '—', '3 listings', '10+ listings'],
        ['Monthly Content Update', '—', '✓ Included', '✓ Included'],
        ['Analytics Dashboard', '—', 'Basic', 'Full'],
        ['Review Management', '—', '—', '✓ Automated'],
        ['Monthly Performance Report', '—', '—', '✓ Included'],
        ['Priority Support', '—', '—', '✓ Included'],
        ['Business Email Hosting', '—', '—', 'Add-on available'],
    ]

    for row_idx, row_data in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = cell_text
            # Format cell
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)

    return table

def main():
    # Create output directory
    output_dir = '/sessions/peaceful-jolly-ritchie/mnt/Website Agency/Templates/Contracts'
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, 'Service_Agreement_v1.0.docx')

    # Create document
    doc = Document()

    # Set up page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('[AGENCY_NAME] — Website Services Agreement')
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.name = 'Arial'
    title.paragraph_format.space_after = Pt(12)

    # Add spacer
    doc.add_paragraph()

    # SECTION 1: PARTIES
    create_heading(doc, 'Section 1: Parties to this Agreement', 1)
    add_body_paragraph(doc, 'This Service Agreement ("Agreement") is entered into as of [DATE] (the "Effective Date") between:')

    add_bullet_point(doc, 'SERVICE PROVIDER: [AGENCY_NAME], operated by Benjamin Rodriguez, a web design and development agency ("Agency")')
    add_bullet_point(doc, 'CLIENT: [CLIENT_NAME], operating as [CLIENT_BUSINESS_NAME] ("Client")')

    add_body_paragraph(doc, 'The Agency and Client are collectively referred to herein as "the Parties."')

    # SECTION 2: SERVICES
    create_heading(doc, 'Section 2: Scope of Services', 1)
    add_body_paragraph(doc, 'The Agency agrees to provide the following website design and development services (the "Services"):')

    add_numbered_item(doc, '2.1', 'Website Design & Development — Professional website design and development using modern, responsive technologies')
    add_numbered_item(doc, '2.2', 'Responsive Mobile Optimization — All websites optimized for mobile devices, tablets, and desktop screens')
    add_numbered_item(doc, '2.3', 'Hosting & Maintenance — Unlimited hosting, security updates, and ongoing maintenance')
    add_numbered_item(doc, '2.4', 'SSL Certificate — Included SSL/TLS encryption for secure data transmission')
    add_numbered_item(doc, '2.5', 'Basic SEO Setup — Initial on-page SEO optimization including meta tags, structured data, and search engine submission')
    add_numbered_item(doc, '2.6', 'Client Revisions — The selected Service Tier includes the specified number of revision rounds')

    add_body_paragraph(doc, 'Additional services, add-ons, and custom development are available at the Agency\'s current rates upon mutual written agreement.')

    # SECTION 3: SERVICE TIERS
    create_heading(doc, 'Section 3: Service Tiers & Pricing', 1)
    add_body_paragraph(doc, 'The Client shall select one of the following Service Tiers:')
    doc.add_paragraph()

    create_service_tiers_table(doc)

    doc.add_paragraph()
    doc.add_paragraph()

    # SECTION 4: PAYMENT TERMS
    create_heading(doc, 'Section 4: Payment Terms & Conditions', 1)

    add_numbered_item(doc, '4.1', 'Setup Fee — The applicable Setup Fee (per selected Tier) is due in full upon execution of this Agreement. Work shall not commence until payment is received.')
    add_numbered_item(doc, '4.2', 'Monthly Service Fee — The Monthly Service Fee (per selected Tier) shall be billed on the [BILLING_DAY]th day of each month. The first month\'s service fee is due within 30 days of invoice.')
    add_numbered_item(doc, '4.3', 'Payment Methods — The Agency accepts payment via: (a) Credit Card (Visa, Mastercard, American Express); (b) ACH Bank Transfer; (c) PayPal; (d) Other agreed methods')
    add_numbered_item(doc, '4.4', 'Late Payment Penalties — Invoices are due net 15 days. If payment is not received within 15 days of invoice date, a late fee of 1.5% per month (18% annually) shall accrue on the outstanding balance.')
    add_numbered_item(doc, '4.5', 'Suspension of Services — If payment is 30 days overdue, the Agency reserves the right to suspend all services, including hosting and maintenance, until payment is received in full.')

    add_body_paragraph(doc, 'All fees are exclusive of applicable sales taxes, which shall be added to invoices if required by law.')

    # SECTION 5: TIMELINE
    create_heading(doc, 'Section 5: Project Timeline & Deliverables', 1)

    add_numbered_item(doc, '5.1', 'Website Draft — The Agency shall deliver an initial website draft within 5 (five) business days of receipt of all required Client information and assets.')
    add_numbered_item(doc, '5.2', 'Client Review Period — The Client shall have 7 (seven) business days to review the website draft and provide feedback.')
    add_numbered_item(doc, '5.3', 'Revisions — The Agency shall complete requested revisions within 3 (three) business days of receipt of Client feedback, subject to the revision limits of the selected Tier.')
    add_numbered_item(doc, '5.4', 'Go-Live — Upon final Client approval, the website shall go live and be accessible at the Client\'s domain within 24 (twenty-four) hours.')

    add_body_paragraph(doc, 'Timelines begin upon Client\'s delivery of all necessary information, content, images, and assets. Client-caused delays extend all timelines accordingly.')

    # SECTION 6: CLIENT RESPONSIBILITIES
    create_heading(doc, 'Section 6: Client Responsibilities', 1)
    add_body_paragraph(doc, 'The Client agrees to:')

    add_numbered_item(doc, '6.1', 'Provide Required Information — Deliver all business information, logos, product/service descriptions, photos, and content within 5 (five) business days of the Effective Date.')
    add_numbered_item(doc, '6.2', 'Timely Feedback — Respond to revision requests and feedback requests within 7 (seven) business days.')
    add_numbered_item(doc, '6.3', 'Accurate Information — Confirm that all information, images, and content provided are accurate, original, or properly licensed.')
    add_numbered_item(doc, '6.4', 'Domain & Hosting Setup — Provide domain access and authorize DNS configuration as required to deploy the website.')

    add_body_paragraph(doc, 'Failure to provide required information or respond to requests shall delay the project timeline proportionately.')

    # SECTION 7: INTELLECTUAL PROPERTY
    create_heading(doc, 'Section 7: Intellectual Property & Ownership', 1)

    add_numbered_item(doc, '7.1', 'Agency Ownership — During the active Service subscription term, the website design, layout, code, and functionality remain the property of [AGENCY_NAME]. The Agency retains the right to showcase the work in its portfolio and case studies.')
    add_numbered_item(doc, '7.2', 'Client Content Ownership — All text, images, logos, and content provided by the Client remain the exclusive property of the Client.')
    add_numbered_item(doc, '7.3', 'Ownership Transfer — Upon Client request and payment of a one-time $497 ownership transfer fee, the Agency shall transfer full ownership of the website code and design to the Client.')
    add_numbered_item(doc, '7.4', 'License Upon Termination — If the Service Agreement terminates without ownership transfer, the Client receives a non-exclusive license to use the published website during the subscription term only.')

    doc.add_paragraph()

    # SECTION 8: TERM & TERMINATION
    create_heading(doc, 'Section 8: Term & Termination', 1)

    add_numbered_item(doc, '8.1', 'Initial Term — This Agreement shall commence on the Effective Date and continue for a minimum of 3 (three) months (the "Initial Term").')
    add_numbered_item(doc, '8.2', 'Month-to-Month — Following the Initial Term, this Agreement shall automatically renew on a month-to-month basis unless either Party provides written notice of termination.')
    add_numbered_item(doc, '8.3', 'Termination Notice — Either Party may terminate this Agreement with 30 (thirty) days\' written notice. The Client remains responsible for all fees during the notice period.')
    add_numbered_item(doc, '8.4', 'Refund Policy — Setup fees are non-refundable after work commences. Monthly service fees are non-refundable but can be credited toward future months if the Client provides 30 days\' notice.')
    add_numbered_item(doc, '8.5', 'Effect of Termination — Upon termination, the Agency shall: (a) cease all services; (b) remove the website from Agency-controlled hosting if ownership was not transferred; (c) provide a copy of the website code if ownership was transferred.')

    doc.add_paragraph()

    # SECTION 9: LIABILITY
    create_heading(doc, 'Section 9: Limitation of Liability & Indemnification', 1)

    add_numbered_item(doc, '9.1', 'Liability Cap — In no event shall either Party\'s total liability exceed the fees paid by the Client in the 3 (three) months preceding the claim.')
    add_numbered_item(doc, '9.2', 'Third-Party Services — The Agency is not liable for outages, downtime, or failures of third-party services including but not limited to: domain registrars, DNS providers, hosting providers, payment processors, or email services.')
    add_numbered_item(doc, '9.3', 'Indemnification — The Client indemnifies the Agency from any claims that Client-provided content infringes third-party intellectual property rights.')
    add_numbered_item(doc, '9.4', 'No Consequential Damages — Neither Party shall be liable for indirect, incidental, consequential, special, or punitive damages.')

    doc.add_paragraph()

    # SECTION 10: MISCELLANEOUS
    create_heading(doc, 'Section 10: Miscellaneous Provisions', 1)

    add_numbered_item(doc, '10.1', 'Governing Law — This Agreement shall be governed by and construed in accordance with the laws of [STATE], without regard to its conflict of law principles.')
    add_numbered_item(doc, '10.2', 'Entire Agreement — This Agreement constitutes the entire agreement between the Parties and supersedes all prior negotiations, understandings, and agreements, whether written or oral.')
    add_numbered_item(doc, '10.3', 'Amendments — No amendment, modification, or waiver of any provision of this Agreement shall be effective unless in writing and signed by both Parties.')
    add_numbered_item(doc, '10.4', 'Severability — If any provision of this Agreement is found to be invalid or unenforceable, such provision shall be modified to the minimum extent necessary to make it enforceable, and the remainder of the Agreement shall remain in effect.')
    add_numbered_item(doc, '10.5', 'Contact Information — All notices required under this Agreement shall be sent to the addresses provided by each Party in writing.')

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # SIGNATURE SECTION
    create_heading(doc, 'Signatures & Acceptance', 1)

    add_body_paragraph(doc, 'By signing below, both Parties acknowledge that they have read, understood, and agree to be bound by the terms of this Service Agreement.')

    doc.add_paragraph()

    # Agency signature
    p = doc.add_paragraph()
    p_run = p.add_run('SERVICE PROVIDER:')
    p_run.font.bold = True
    p_run.font.name = 'Arial'
    p_run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph('_________________________________________________')
    doc.add_paragraph('[AGENCY_NAME]')
    doc.add_paragraph('Authorized Representative: Benjamin Rodriguez')
    doc.add_paragraph('Date: _____________________')

    doc.add_paragraph()
    doc.add_paragraph()

    # Client signature
    p = doc.add_paragraph()
    p_run = p.add_run('CLIENT:')
    p_run.font.bold = True
    p_run.font.name = 'Arial'
    p_run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph('_________________________________________________')
    doc.add_paragraph('Client Name (Print): _________________________________')
    doc.add_paragraph('Business Name (Print): _______________________________')
    doc.add_paragraph('Title/Role: __________________________________________')
    doc.add_paragraph('Date: _____________________')

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer signature
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run = p.add_run('— Benjamin Rodriguez')
    p_run.font.italic = True
    p_run.font.name = 'Arial'
    p_run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(12)

    # Save document
    doc.save(output_path)

    # Get file info
    file_size = os.path.getsize(output_path)

    print(f'✓ Service Agreement contract generated successfully')
    print(f'✓ Saved to: {output_path}')
    print(f'✓ File size: {file_size / 1024:.2f} KB')

    return 0

if __name__ == '__main__':
    sys.exit(main())
